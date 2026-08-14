"""Reconstruct files with pseudonymised content.

Only a subset of formats support in-place reconstruction:

* **DOCX** – replace text in paragraph runs via ``python-docx``
* **XLSX** – replace cell values via ``openpyxl``
* **Plain text** – simple string encode

Formats like PDF and images *cannot* be reconstructed; the caller should
convert those blocks to text instead.

**The guarantee.** A reconstructed file is shipped only if extracting it again
yields the masked text — otherwise the caller converts the file to text
instead. Replacement reaches paragraph runs and string cells; a document has
surfaces that are neither (hyperlink runs, numeric cells), and a detected span
can cross a paragraph break that no single paragraph contains. Reasoning about
the map alone cannot see any of that, so the artifact is verified rather than
the plan. The failure mode is a file that loses its formatting, never one that
ships PII the masked text had removed.

The only accepted difference between the two is a placeholder that replaces
nothing: a zero-length span mints one, no text replacement can place it, and
it masks nothing.
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING, Protocol

import structlog

if TYPE_CHECKING:
    from collections.abc import Callable

    from docx.table import Table
    from docx.text.paragraph import Paragraph

    from noirdoc.file_analysis.models import FileBlock
    from noirdoc.pseudonymization.engine import PseudonymizationEngine
    from noirdoc.pseudonymization.mapper import PseudonymMapper

logger = structlog.get_logger()


class _BlockContainer(Protocol):
    """Common interface of python-docx block containers (body, header, footer, cell)."""

    @property
    def paragraphs(self) -> list[Paragraph]: ...

    @property
    def tables(self) -> list[Table]: ...


_RECONSTRUCTABLE_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/plain",
    "text/csv",
    "text/markdown",
    "text/html",
}


def can_reconstruct(mime_type: str) -> bool:
    """Return ``True`` if in-place content replacement is supported for *mime_type*."""
    return mime_type in _RECONSTRUCTABLE_MIMES


def pseudonymize_block(
    block: FileBlock,
    mapper: PseudonymMapper,
    engine: PseudonymizationEngine | None = None,
) -> str | None:
    """Pseudonymize ``block.extracted_text`` and record what reconstruction needs.

    Sets ``pseudonymized_text`` and ``emitted_spans`` together. Reconstruction
    rewrites the file itself — DOCX runs, XLSX cells — not the extracted text,
    so it needs to know which placeholder replaced which original. Producing
    both here keeps that pairing with the substitution that created it;
    reconstructing it afterwards from entity offsets means a second copy of
    the overlap semantics, which is exactly what used to drift.
    """
    from noirdoc.pseudonymization.engine import PseudonymizationEngine

    if block.extracted_text is None:
        return None

    text, spans = (engine or PseudonymizationEngine()).pseudonymize_detailed(
        block.extracted_text,
        block.entities,
        mapper,
    )
    block.pseudonymized_text = text
    block.emitted_spans = spans
    return text


def reconstruct(block: FileBlock) -> bytes | None:
    """Return new file bytes with pseudonymised text, or ``None`` on failure.

    The caller must ensure ``block.pseudonymized_text`` is set and
    ``can_reconstruct(block.mime_type)`` is ``True``.

    Bytes that come out are kept on the block, so a later reader — the
    pipeline's accounting, a second apply pass — can tell a file that was
    rebuilt from one that was refused without rebuilding it again.
    """
    # Short-circuit: pre-computed bytes (e.g. from XLSX column-inference)
    if block.reconstructed_bytes is not None:
        return block.reconstructed_bytes

    if block.pseudonymized_text is None:
        return None

    mime = block.mime_type
    data: bytes | None = None
    if mime in ("text/plain", "text/csv", "text/markdown", "text/html"):
        data = _reconstruct_plain(block)
    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        data = _reconstruct_docx(block)
    elif mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        data = _reconstruct_xlsx(block)

    if data is not None:
        block.reconstructed_bytes = data
    return data


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


def _reconstruct_plain(block: FileBlock) -> bytes:
    return (block.pseudonymized_text or "").encode("utf-8")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _replace_in_block_container(container: _BlockContainer, replacements: dict[str, str]) -> None:
    """Apply *replacements* to every paragraph in *container* and its tables."""
    for para in container.paragraphs:
        _replace_in_paragraph(para, replacements)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)


def _reconstruct_docx(block: FileBlock) -> bytes | None:
    """Find-and-replace detected entities in DOCX paragraph runs.

    Covers the document body, every section's headers and footers
    (default, first-page, even-page variants), and review comments —
    matching the surfaces walked by :func:`extract_docx` so PII the
    detector saw is also stripped from the output bytes. What it writes is
    verified by re-extracting it; see the module docstring.
    """
    from docx import Document

    from noirdoc.file_analysis.extractors.docx_ext import extract_docx

    try:
        doc = Document(io.BytesIO(block.content_bytes))
    except Exception:
        return None

    replacements = _build_replacements(block)
    if replacements is None:
        return None  # cannot pair placeholders with originals — refuse
    if not replacements:
        return block.content_bytes  # nothing to replace

    _replace_in_block_container(doc, replacements)

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            _replace_in_block_container(header, replacements)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            _replace_in_block_container(footer, replacements)

    try:
        comments = list(doc.comments)
    except Exception:
        comments = []
    for comment in comments:
        _replace_in_block_container(comment, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if not _written_file_matches_the_masked_text(block, data, extract_docx):
        return None
    return data


def _replace_in_paragraph(para: Paragraph, replacements: dict[str, str]) -> None:
    """Replace entity text across runs in a paragraph."""
    full_text = para.text
    for original, pseudo in replacements.items():
        if original in full_text:
            full_text = full_text.replace(original, pseudo)

    if full_text == para.text:
        return

    # Re-write runs: clear all but first, set first to full text.
    # This simplifies formatting but is acceptable for v1.
    if para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ""


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def _reconstruct_xlsx(block: FileBlock) -> bytes | None:
    """Replace cell values that contain detected entities.

    Only string cells are rewritten, while the extractor stringifies every
    cell — a number the detector flagged would survive. What is written is
    verified by re-extracting it; see the module docstring.
    """
    from openpyxl import load_workbook

    from noirdoc.file_analysis.extractors.xlsx import extract_xlsx

    try:
        wb = load_workbook(io.BytesIO(block.content_bytes))
    except Exception:
        return None

    replacements = _build_replacements(block)
    if replacements is None:
        wb.close()
        return None  # cannot pair placeholders with originals — refuse
    if not replacements:
        buf = io.BytesIO()
        wb.save(buf)
        wb.close()
        return buf.getvalue()

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    new_val = cell.value
                    for original, pseudo in replacements.items():
                        new_val = new_val.replace(original, pseudo)
                    if new_val != cell.value:
                        cell.value = new_val

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    data = buf.getvalue()
    if not _written_file_matches_the_masked_text(block, data, extract_xlsx):
        return None
    return data


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _build_replacements(block: FileBlock) -> dict[str, str] | None:
    """Map original text -> placeholder from what the pseudonymizer emitted.

    The pairing comes from :func:`pseudonymize_block`, which records it while
    substituting. It is not re-derived from entity offsets: an overlapping
    entity contributes a placeholder of its own for the remainder that reaches
    past the winning span, and offset arithmetic over the entity list cannot
    see it — every later entity then lands on the wrong offset and is either
    dropped or paired with someone else's placeholder.

    Longest original first, so a short original cannot eat a longer one it is
    contained in ("Weber" inside "Anna Weber") before that longer one is
    replaced. A span with an empty original is left out: it replaces no
    characters, and ``str.replace("", …)`` would splice its placeholder
    between every character in the document.

    Returns ``None`` when the block has entities but no record of the
    substitution, and when the map cannot reproduce the masked text (see
    :func:`_replacements_reproduce_the_masked_text`): the caller must not
    rewrite the file from a guess, and falling back to the converted text is
    the safe answer.
    """
    if not block.entities or not block.extracted_text or not block.pseudonymized_text:
        return {}

    if block.emitted_spans is None:
        logger.warning(
            "reconstruction.missing_emitted_spans",
            path=block.source_path,
            mime=block.mime_type,
            entity_count=len(block.entities),
        )
        return None

    spans = sorted(block.emitted_spans, key=lambda s: len(s.original), reverse=True)
    replacements = {span.original: span.pseudonym for span in spans if span.original}

    if not _replacements_reproduce_the_masked_text(replacements, block):
        return None
    return replacements


def _replacements_reproduce_the_masked_text(
    replacements: dict[str, str],
    block: FileBlock,
) -> bool:
    """Check the map against the text the pseudonymizer actually produced.

    A cheap early-out on everything that goes wrong in the map itself, before
    a file is parsed and rewritten. Replacement rewrites by text, not by
    offset, so an original that also occurs inside an already-inserted
    placeholder is rewritten there too — an ID "12" alongside twelve people
    turns ``<<PERSON_12>>`` into ``<<PERSON_<<ID_1>>>>``, still masked but no
    longer revealable — and an original can equally hit an occurrence the
    detector never flagged ("Weber" inside "Weberstrasse"). Both show up as a
    rebuild that no longer matches.

    It does NOT say the file will come out right: whether the writers can
    reach every place the text lives is a property of the file, checked after
    writing by :func:`_written_file_matches_the_masked_text`.

    An original inside its OWN placeholder is not a hazard and must not be
    refused: ``str.replace`` makes one left-to-right pass and never rescans
    what it inserted, and a shorter original's placeholder does not exist in
    the text yet while the longer originals are being applied.

    Nothing about the offending values is logged — they are the PII this
    module exists to remove.
    """
    rebuilt = block.extracted_text or ""
    for original, pseudonym in replacements.items():
        rebuilt = rebuilt.replace(original, pseudonym)
    if rebuilt != _reference_text(block):
        logger.warning(
            "reconstruction.replacement_selfcheck_failed",
            path=block.source_path,
            mime=block.mime_type,
            replacement_count=len(replacements),
        )
        return False
    return True


def _reference_text(block: FileBlock) -> str:
    """The masked text as a reconstructed file can carry it.

    Identical to ``pseudonymized_text`` except for placeholders that replace
    nothing: a zero-length span mints one, and no text replacement can put it
    into a file. Leaving it out costs an empty annotation — its mapping's
    original is the empty string, so no PII rides on it.
    """
    text = block.pseudonymized_text or ""
    for span in block.emitted_spans or ():
        if not span.original:
            text = text.replace(span.pseudonym, "")
    return text


def _written_file_matches_the_masked_text(
    block: FileBlock,
    data: bytes,
    extract: Callable[[bytes], str],
) -> bool:
    """Read the file that was just written and require the masked text back.

    The replacement map can be perfect and the write still incomplete. The
    DOCX writer rewrites ``paragraph.runs``, which excludes hyperlink runs
    that ``paragraph.text`` includes; a detected span can cross the paragraph
    break the extractor joins with a newline and match no paragraph at all;
    the XLSX writer only touches string cells while the extractor stringifies
    numbers too. Each of those ships the original beside its own placeholder,
    and none of them is visible in the map.

    Extracting the written bytes with the same extractor that produced the
    input text is what sees them. On a mismatch the file is refused and the
    caller converts to text: formatting is lost, masking is not.
    """
    try:
        rebuilt = extract(data)
    except Exception:
        rebuilt = None
    if rebuilt == _reference_text(block):
        return True
    logger.warning(
        "reconstruction.postwrite_verification_failed",
        path=block.source_path,
        mime=block.mime_type,
    )
    return False

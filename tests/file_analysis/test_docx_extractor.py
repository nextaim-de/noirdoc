"""NDS-016: DOCX extraction must cover header/footer/comment surfaces."""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

from noirdoc.detection.base import DetectedEntity
from noirdoc.file_analysis.extractors.docx_ext import extract_docx
from noirdoc.file_analysis.models import FileBlock
from noirdoc.file_analysis.reconstruction import _reconstruct_docx

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph


def _entity(text: str, in_text: str) -> DetectedEntity:
    start = in_text.index(text)
    return DetectedEntity(
        entity_type="PERSON",
        text=text,
        start=start,
        end=start + len(text),
        score=0.9,
        source="test",
    )


def _add_hyperlink(para: Paragraph, text: str, url: str) -> None:
    """Append a real ``w:hyperlink`` (with one run) to *para*."""
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    r_id = para.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


def _docx_block(docx_bytes: bytes, pseudo_by_original: dict[str, str]) -> FileBlock:
    """Build a FileBlock with synthetic pseudonymized text for *docx_bytes*."""
    extracted = extract_docx(docx_bytes)
    pseudonymized = extracted
    entities = []
    for original, pseudo in pseudo_by_original.items():
        pseudonymized = pseudonymized.replace(original, pseudo)
        entities.append(_entity(original, extracted))
    return FileBlock(
        content_bytes=docx_bytes,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source_path="test.docx",
        source_type="file",
        extracted_text=extracted,
        pseudonymized_text=pseudonymized,
        entities=entities,
    )


def _docx_with_headers_footers_and_body() -> bytes:
    """Build a DOCX whose header, footer, and body each contain distinct PII."""
    from docx import Document

    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Header: Anna Mueller"
    section.footer.paragraphs[0].text = "Footer: Bernd Schmidt"
    doc.add_paragraph("Body: Carla Weber")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx_walks_headers_and_footers():
    """PII embedded in section headers and footers must reach the detector."""
    text = extract_docx(_docx_with_headers_footers_and_body())
    assert "Anna Mueller" in text
    assert "Bernd Schmidt" in text
    assert "Carla Weber" in text


def test_extract_docx_walks_comments():
    """Review comments are a routine PII surface — they must be extracted."""
    from docx import Document

    doc = Document()
    para = doc.add_paragraph("Body text")
    doc.add_comment(runs=[para.runs[0]] if para.runs else [], text="Reviewer: Dora Klein")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_docx(buf.getvalue())
    assert "Dora Klein" in text


def test_reconstruct_docx_replaces_text_in_headers_and_footers():
    """Reconstruction must scrub header/footer text so the output bytes are clean."""
    docx_bytes = _docx_with_headers_footers_and_body()
    extracted = extract_docx(docx_bytes)

    # Build a synthetic pseudonymized result: replace each name with a token
    pseudonymized = (
        extracted.replace("Anna Mueller", "<<PERSON_1>>")
        .replace("Bernd Schmidt", "<<PERSON_2>>")
        .replace("Carla Weber", "<<PERSON_3>>")
    )

    entities = [
        _entity("Anna Mueller", extracted),
        _entity("Bernd Schmidt", extracted),
        _entity("Carla Weber", extracted),
    ]

    block = FileBlock(
        content_bytes=docx_bytes,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source_path="test.docx",
        source_type="file",
        extracted_text=extracted,
        pseudonymized_text=pseudonymized,
        entities=entities,
    )

    new_bytes = _reconstruct_docx(block)
    assert new_bytes is not None

    rewritten = extract_docx(new_bytes)
    assert "Anna Mueller" not in rewritten
    assert "Bernd Schmidt" not in rewritten
    assert "Carla Weber" not in rewritten
    assert "<<PERSON_1>>" in rewritten
    assert "<<PERSON_2>>" in rewritten
    assert "<<PERSON_3>>" in rewritten


def test_reconstruct_docx_rewrites_entity_inside_hyperlink():
    """Issue #16 shape 1: an entity inside a hyperlink must be replaced, not leaked.

    Before the fix the placeholder was prepended to the paragraph while the
    original address stayed in the hyperlink run — a leak in shipped bytes.
    """
    from docx import Document

    doc = Document()
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.add_run("Kontakt: ")
    _add_hyperlink(footer_para, "info@musterfirma.de", "mailto:info@musterfirma.de")
    buf = io.BytesIO()
    doc.save(buf)

    block = _docx_block(buf.getvalue(), {"info@musterfirma.de": "<<EMAIL_1>>"})
    new_bytes = _reconstruct_docx(block)
    assert new_bytes is not None

    result = Document(io.BytesIO(new_bytes))
    para = result.sections[0].footer.paragraphs[0]
    assert para.text == "Kontakt: <<EMAIL_1>>"
    assert "info@musterfirma.de" not in para.text
    # The placeholder must live inside the hyperlink itself
    assert para.hyperlinks[0].text == "<<EMAIL_1>>"


def test_reconstruct_docx_does_not_duplicate_unrelated_hyperlink():
    """Issue #16 shape 2: a hyperlink in the same paragraph must not be copied into runs."""
    from docx import Document

    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Anna ")
    para.add_run("Weber (siehe ")
    _add_hyperlink(para, "www.firma.de", "https://www.firma.de")
    para.add_run(")")
    buf = io.BytesIO()
    doc.save(buf)

    block = _docx_block(buf.getvalue(), {"Anna Weber": "<<PERSON_1>>"})
    new_bytes = _reconstruct_docx(block)
    assert new_bytes is not None

    result = Document(io.BytesIO(new_bytes))
    out = result.paragraphs[0]
    assert out.text == "<<PERSON_1>> (siehe www.firma.de)"
    # Hyperlink text stays in the hyperlink, never in the plain runs
    assert out.hyperlinks[0].text == "www.firma.de"
    plain_runs_text = "".join(run.text for run in out.runs)
    assert "www.firma.de" not in plain_runs_text

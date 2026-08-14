"""DOCX/XLSX reconstruction must replace exactly what pseudonymize replaced.

Reconstruction is a find-and-replace over the *file* (paragraph runs, cell
values) rather than over the extracted text, so it needs one thing the
extracted text cannot tell it: which placeholder stands for which original.
Deriving that from entity offsets is a second implementation of the builder's
overlap semantics, and it drifts — an overlapping entity contributes a
remainder pseudonym that entity-offset arithmetic cannot see, and every later
entity is then looked up at the wrong offset.

These tests pin the contract end to end: every placeholder in the
pseudonymized text is in the replacement map under the exact original it
replaced, applying the map to the extracted text reproduces the pseudonymized
text byte for byte, and — the only claim that covers what the writers can and
cannot reach — re-extracting the file that was written reproduces it too.
"""

from __future__ import annotations

import io
import random
import re
import string

from structlog.testing import capture_logs

from noirdoc.detection.base import DetectedEntity
from noirdoc.file_analysis.models import FileBlock
from noirdoc.file_analysis.reconstruction import (
    _build_replacements,
    _reconstruct_docx,
    _reconstruct_xlsx,
    pseudonymize_block,
)
from noirdoc.pseudonymization.mapper import PseudonymMapper

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

_PLACEHOLDER = re.compile(r"<<[A-Z_]+_\d+>>")


def _ent(entity_type: str, text: str, start: int, end: int) -> DetectedEntity:
    return DetectedEntity(
        entity_type=entity_type,
        text=text,
        start=start,
        end=end,
        score=0.9,
        source="test",
    )


def _block(mime: str, content: bytes, text: str, entities: list[DetectedEntity]) -> FileBlock:
    """Build a block the way the pipeline does: extract, detect, pseudonymize."""
    block = FileBlock(
        content_bytes=content,
        mime_type=mime,
        source_path="test",
        source_type="file",
        extracted_text=text,
        entities=entities,
    )
    return block


def _docx_bytes(*paragraphs: str) -> bytes:
    from docx import Document

    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_with_hyperlink() -> bytes:
    """A paragraph whose second half lives in a hyperlink run.

    python-docx counts hyperlink text in `Paragraph.text` but leaves it out of
    `Paragraph.runs`, so the extractor sees it and the writer cannot touch it.
    """
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Document()
    para = doc.add_paragraph("Kontakt: ")
    para._p.append(
        parse_xml(
            f"<w:hyperlink {nsdecls('w')}>"
            f'<w:r><w:t xml:space="preserve">Anna Beispiel</w:t></w:r>'
            f"</w:hyperlink>"
        )
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_text(data: bytes) -> str:
    """Read the file back the way the pipeline reads it."""
    from noirdoc.file_analysis.extractors.docx_ext import extract_docx

    return extract_docx(data)


def _xlsx_bytes(*cell_values: object) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    for row, value in enumerate(cell_values, start=1):
        wb.active[f"A{row}"] = value
    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    return buf.getvalue()


def _xlsx_text(data: bytes) -> str:
    """Read the file back the way the pipeline reads it."""
    from noirdoc.file_analysis.extractors.xlsx import extract_xlsx

    return extract_xlsx(data)


def _apply(replacements: dict[str, str], text: str) -> str:
    """Apply the replacement map the way the DOCX/XLSX writers do."""
    for original, pseudonym in replacements.items():
        text = text.replace(original, pseudonym)
    return text


# --- The mixed document: one overlapping pair among plain entities ---------

_MIXED = "Anna Weber wohnt in Berlin und arbeitet bei Weber GmbH in Hamburg."


def _mixed_entities() -> list[DetectedEntity]:
    org = _MIXED.index("Weber GmbH")
    berlin = _MIXED.index("Berlin")
    hamburg = _MIXED.index("Hamburg")
    return [
        _ent("PERSON", "Anna Weber", 0, 10),
        _ent("LOCATION", "Berlin", berlin, berlin + 6),
        _ent("PERSON", "Weber", org, org + 5),
        _ent("ORGANIZATION", "Weber GmbH", org, org + 10),
        _ent("LOCATION", "Hamburg", hamburg, hamburg + 7),
    ]


def test_replacements_cover_every_placeholder_of_an_overlapping_document():
    """The remainder pseudonym must not knock the later entities off their offsets."""
    mapper = PseudonymMapper()
    block = _block(_DOCX_MIME, _docx_bytes(_MIXED), _MIXED, _mixed_entities())
    pseudonymize_block(block, mapper)

    replacements = _build_replacements(block)

    assert replacements == {
        "Anna Weber": "<<PERSON_2>>",
        "Hamburg": "<<LOCATION_1>>",
        "Berlin": "<<LOCATION_2>>",
        "Weber": "<<PERSON_1>>",
        "GmbH": "<<ORGANIZATION_1>>",
    }
    # Longest original first, so "Weber" cannot eat the "Weber" inside "Anna Weber".
    assert list(replacements) == sorted(replacements, key=len, reverse=True)
    assert _apply(replacements, _MIXED) == block.pseudonymized_text


def test_reconstruct_docx_matches_the_pseudonymized_text_byte_for_byte():
    mapper = PseudonymMapper()
    block = _block(_DOCX_MIME, _docx_bytes(_MIXED), _MIXED, _mixed_entities())
    pseudonymize_block(block, mapper)

    new_bytes = _reconstruct_docx(block)
    assert new_bytes is not None
    assert _docx_text(new_bytes) == block.pseudonymized_text
    assert "GmbH" not in _docx_text(new_bytes)
    assert "Hamburg" not in _docx_text(new_bytes)


def test_reconstruct_xlsx_matches_the_pseudonymized_text_byte_for_byte():
    mapper = PseudonymMapper()
    block = _block(_XLSX_MIME, _xlsx_bytes(_MIXED), _MIXED, _mixed_entities())
    pseudonymize_block(block, mapper)

    new_bytes = _reconstruct_xlsx(block)
    assert new_bytes is not None
    assert _xlsx_text(new_bytes) == block.pseudonymized_text
    assert "GmbH" not in _xlsx_text(new_bytes)
    assert "Hamburg" not in _xlsx_text(new_bytes)


# --- The replacement map has to reproduce the masked text -----------------


def test_replacements_refused_when_an_original_hides_in_a_placeholder():
    """An original made of placeholder characters can corrupt a placeholder.

    Replacement rewrites the file by text, so an original that also occurs
    INSIDE an already-inserted placeholder gets rewritten there too. "12" is
    an ID here and also the counter of `<<PERSON_12>>`, which would come out
    as `<<PERSON_<<ID_1>>>>` — masked, but no longer revealable.

    Reconstruction refuses instead: the caller falls back to converted text,
    which keeps both the masking and the reveal and only loses the formatting.
    """
    people = [f"N{i:02d}" for i in range(12)]
    text = "code 12 dann " + " ".join(people)
    entities = [_ent("ID", "12", text.index("12"), text.index("12") + 2)]
    for name in people:
        start = text.index(name)
        entities.append(_ent("PERSON", name, start, start + len(name)))

    mapper = PseudonymMapper()
    block = _block(_DOCX_MIME, _docx_bytes(text), text, entities)
    pseudonymize_block(block, mapper)
    assert "<<PERSON_12>>" in (block.pseudonymized_text or "")

    with capture_logs() as logs:
        assert _build_replacements(block) is None
    assert [entry["event"] for entry in logs] == ["reconstruction.replacement_selfcheck_failed"]

    assert _reconstruct_docx(block) is None
    assert _reconstruct_xlsx(block) is None


def test_replacements_allow_an_original_inside_its_own_placeholder():
    """An original that occurs in its OWN placeholder is benign — do not refuse it.

    An ID "1" mints `<<ID_1>>`, which contains "1". Replacement still lands
    correctly: `str.replace` makes one left-to-right pass and never rescans
    what it inserted. Refusing this shape would drop the formatting of an
    ordinary document for nothing.
    """
    text = "Der Vorgang 1 wurde geprueft."
    entities = [_ent("ID", "1", text.index("1"), text.index("1") + 1)]

    mapper = PseudonymMapper()
    block = _block(_DOCX_MIME, _docx_bytes(text), text, entities)
    pseudonymize_block(block, mapper)
    assert block.pseudonymized_text == "Der Vorgang <<ID_1>> wurde geprueft."

    with capture_logs() as logs:
        replacements = _build_replacements(block)
    assert replacements == {"1": "<<ID_1>>"}
    assert logs == []
    assert _docx_text(_reconstruct_docx(block) or b"") == block.pseudonymized_text


def test_replacements_pass_the_self_check_on_a_normal_document():
    """The guard must not fire on ordinary text — reconstruction proceeds."""
    mapper = PseudonymMapper()
    block = _block(_DOCX_MIME, _docx_bytes(_MIXED), _MIXED, _mixed_entities())
    pseudonymize_block(block, mapper)

    with capture_logs() as logs:
        replacements = _build_replacements(block)
    assert replacements is not None
    assert logs == []
    assert _apply(replacements, _MIXED) == block.pseudonymized_text
    assert _reconstruct_docx(block) is not None


# --- The file that gets shipped has to say what the masked text says ------
#
# The replacement map can be perfect and the write still incomplete: the
# writers reach paragraph runs and string cells, and a document has surfaces
# that are neither. Re-extracting the bytes that were written is the only
# check that sees those.


def test_docx_hyperlink_run_leak_is_refused():
    """PII in a hyperlink run is extracted but cannot be written — refuse.

    python-docx puts hyperlink text in `Paragraph.text` but not in
    `Paragraph.runs`. The detector sees "Anna Beispiel", the map contains it,
    and the writer — which rewrites `runs[0]` and blanks the rest — cannot
    touch it: the shipped paragraph would read "Kontakt: <<PERSON_1>>Anna
    Beispiel", the name masked and duplicated in cleartext beside it.
    """
    content = _docx_with_hyperlink()
    text = _docx_text(content)
    assert text == "Kontakt: Anna Beispiel"
    start = text.index("Anna Beispiel")
    entities = [_ent("PERSON", "Anna Beispiel", start, start + len("Anna Beispiel"))]

    block = _block(_DOCX_MIME, content, text, entities)
    pseudonymize_block(block, PseudonymMapper())

    with capture_logs() as logs:
        result = _reconstruct_docx(block)
    assert result is None
    assert [entry["event"] for entry in logs] == ["reconstruction.postwrite_verification_failed"]


def test_docx_entity_spanning_two_paragraphs_is_refused():
    """An entity across a paragraph break matches the flat text and no paragraph.

    `extract_docx` joins paragraphs with "\\n", so a detector can flag a span
    that no single paragraph contains. The map is consistent with the flat
    text — the pre-apply check passes — but the writer never finds the
    original, and the organisation would ship in cleartext.
    """
    content = _docx_bytes("Rechnung von Muster", "GmbH & Co in Hamburg")
    text = _docx_text(content)
    assert text == "Rechnung von Muster\nGmbH & Co in Hamburg"
    start = text.index("Muster")
    entities = [_ent("ORGANIZATION", "Muster\nGmbH & Co", start, start + len("Muster\nGmbH & Co"))]

    block = _block(_DOCX_MIME, content, text, entities)
    pseudonymize_block(block, PseudonymMapper())
    # The map itself is fine — this is not something the pre-apply check sees.
    assert _apply(_build_replacements(block) or {}, text) == block.pseudonymized_text

    with capture_logs() as logs:
        result = _reconstruct_docx(block)
    assert result is None
    assert [entry["event"] for entry in logs] == ["reconstruction.postwrite_verification_failed"]


def test_xlsx_numeric_cell_is_refused():
    """The XLSX writer only rewrites string cells; the extractor stringifies all.

    A customer number stored as a number is extracted as "480815", detected,
    and masked in the text — but `cell.value` is an int, the writer skips it,
    and the number ships unchanged.
    """
    content = _xlsx_bytes(480815, "Anna Weber")
    text = _xlsx_text(content)
    assert text == "480815\nAnna Weber"
    entities = [
        _ent("ID", "480815", 0, 6),
        _ent("PERSON", "Anna Weber", 7, 17),
    ]

    block = _block(_XLSX_MIME, content, text, entities)
    pseudonymize_block(block, PseudonymMapper())

    with capture_logs() as logs:
        result = _reconstruct_xlsx(block)
    assert result is None
    assert [entry["event"] for entry in logs] == ["reconstruction.postwrite_verification_failed"]


def test_over_masking_collision_is_refused():
    """Replacement is by text, so an unflagged occurrence would be masked too.

    Only the standalone "Weber" is an entity; the "Weber" inside
    "Weberstrasse" is not. Replacing by text would mask both, which the
    pre-apply check catches — the document costs its formatting, never its
    masking.
    """
    text = "Weber wohnt in der Weberstrasse."
    entities = [_ent("PERSON", "Weber", 0, 5)]

    block = _block(_DOCX_MIME, _docx_bytes(text), text, entities)
    pseudonymize_block(block, PseudonymMapper())

    with capture_logs() as logs:
        assert _build_replacements(block) is None
    assert [entry["event"] for entry in logs] == ["reconstruction.replacement_selfcheck_failed"]
    assert _reconstruct_docx(block) is None


def test_multi_paragraph_document_passes_post_write_verification():
    """The verification must not fire on an ordinary multi-paragraph document."""
    content = _docx_bytes("Anna Weber wohnt in Berlin.", "Kontakt: max@test.de")
    text = _docx_text(content)
    entities = [
        _ent("PERSON", "Anna Weber", 0, 10),
        _ent("LOCATION", "Berlin", text.index("Berlin"), text.index("Berlin") + 6),
        _ent(
            "EMAIL",
            "max@test.de",
            text.index("max@test.de"),
            text.index("max@test.de") + 11,
        ),
    ]

    block = _block(_DOCX_MIME, content, text, entities)
    pseudonymize_block(block, PseudonymMapper())

    with capture_logs() as logs:
        new_bytes = _reconstruct_docx(block)
    assert new_bytes is not None
    assert logs == []
    assert _docx_text(new_bytes) == block.pseudonymized_text


def test_zero_length_entity_is_left_out_of_the_replacement_map():
    """A placeholder that replaces nothing has nothing to anchor it in the file.

    A zero-length span mints a pseudonym whose original is "" — masking no
    characters at all. Feeding "" to `str.replace` would splice the pseudonym
    between every character in the document, so it is left out of the map, and
    the file ships without that (empty) annotation. No PII rides on it: the
    mapping's original is the empty string.
    """
    text = "Anna Weber wohnt in Berlin."
    entities = [
        _ent("PERSON", "Anna Weber", 0, 10),
        _ent("LOCATION", "", 20, 20),
    ]

    block = _block(_DOCX_MIME, _docx_bytes(text), text, entities)
    pseudonymize_block(block, PseudonymMapper())
    assert block.pseudonymized_text == "<<PERSON_1>> wohnt in <<LOCATION_1>>Berlin."

    with capture_logs() as logs:
        replacements = _build_replacements(block)
        new_bytes = _reconstruct_docx(block)
    assert replacements == {"Anna Weber": "<<PERSON_1>>"}
    assert logs == []
    assert new_bytes is not None
    assert _docx_text(new_bytes) == "<<PERSON_1>> wohnt in Berlin."


# --- No overlap: the behaviour that already worked must keep working -------

_PLAIN = "Anna Weber wohnt in Berlin und schreibt an max@test.de."


def test_reconstruct_without_overlaps_is_unchanged():
    """The pre-existing (non-overlapping) case: same replacements, same output."""
    mapper = PseudonymMapper()
    email = _PLAIN.index("max@test.de")
    berlin = _PLAIN.index("Berlin")
    entities = [
        _ent("PERSON", "Anna Weber", 0, 10),
        _ent("LOCATION", "Berlin", berlin, berlin + 6),
        _ent("EMAIL", "max@test.de", email, email + 11),
    ]
    block = _block(_DOCX_MIME, _docx_bytes(_PLAIN), _PLAIN, entities)
    pseudonymize_block(block, mapper)

    assert _build_replacements(block) == {
        "Anna Weber": "<<PERSON_1>>",
        "max@test.de": "<<EMAIL_1>>",
        "Berlin": "<<LOCATION_1>>",
    }
    new_bytes = _reconstruct_docx(block)
    assert new_bytes is not None
    assert _docx_text(new_bytes) == block.pseudonymized_text


# --- Randomized sweep over overlapping shapes -----------------------------
#
# Same alphabet discipline as the engine's property test: unique
# non-whitespace characters, disjoint from the placeholder charset, so every
# span text occurs exactly once and a replacement can be checked by substring.
# Entity edges are snapped off whitespace, the way real detector spans are.
#
# That disjointness also makes this sweep blind to originals drawn from the
# placeholder charset itself ("12" inside `<<PERSON_12>>`) — deliberately, so
# the invariants stay decidable. `..._an_original_hides_in_a_placeholder`
# above covers that class.

_UNIQUE_ALPHABET = string.ascii_lowercase + "!#$%&()*+,-./:;=?@[]^{|}~"
_TYPES = ("PERSON", "LOCATION", "EMAIL", "ORGANIZATION")


def _random_text(rng: random.Random) -> str:
    out: list[str] = []
    for char in _UNIQUE_ALPHABET[: rng.randrange(20, len(_UNIQUE_ALPHABET) + 1)]:
        out.append(char)
        if rng.random() < 0.2:
            out.append(" ")
    return "".join(out)


def _random_entities(rng: random.Random, text: str) -> list[DetectedEntity]:
    entities: list[DetectedEntity] = []
    for _ in range(rng.randrange(1, 9)):
        start = rng.randrange(len(text))
        end = min(len(text), start + rng.randrange(1, 12))
        # Detector spans never carry edge whitespace.
        while start < end and text[start].isspace():
            start += 1
        while end > start and text[end - 1].isspace():
            end -= 1
        if start == end:
            continue
        entities.append(_ent(rng.choice(_TYPES), text[start:end], start, end))
    rng.shuffle(entities)
    return entities


def test_replacement_sweep_over_overlapping_documents():
    """Zero missing and zero wrong replacements across 3 000 randomized cases."""
    rng = random.Random(20260814)
    missing = 0
    wrong = 0
    not_reproduced = 0
    overlapping_cases = 0

    for _ in range(3000):
        text = _random_text(rng)
        entities = _random_entities(rng, text)
        if not entities:
            continue

        mapper = PseudonymMapper()
        block = _block(_DOCX_MIME, b"", text, entities)
        pseudonymize_block(block, mapper)
        pseudonymized = block.pseudonymized_text or ""

        # Expected map, derived from the output alone: every placeholder in it,
        # under the original the mapper hands back for it.
        expected = {
            mapper.reverse_lookup(token): token for token in _PLACEHOLDER.findall(pseudonymized)
        }
        entity_texts = {e.text for e in entities}
        if any(span.original not in entity_texts for span in block.emitted_spans or []):
            overlapping_cases += 1

        replacements = _build_replacements(block) or {}
        for original, token in expected.items():
            assert original is not None
            if original not in replacements:
                missing += 1
            elif replacements[original] != token:
                wrong += 1

        if _apply(replacements, text) != pseudonymized:
            not_reproduced += 1

    assert (missing, wrong, not_reproduced) == (0, 0, 0)
    # Guard against a vacuous run: remainders must actually have occurred.
    assert overlapping_cases > 300

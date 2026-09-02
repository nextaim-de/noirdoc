"""DOCX package parts outside the body must be pseudonymized and revealed.

Covers ``docProps/core.xml``, ``docProps/app.xml``, ``docProps/custom.xml``,
comment authors/initials, ``word/people.xml``, the ``docProps/thumbnail.jpeg``
preview and ``customXml/`` islands — every surface a plain python-docx
round-trip preserves verbatim (and that a redacted document therefore leaked).
"""

from __future__ import annotations

import io

from docx import Document
from docx.document import Document as DocumentObject

from noirdoc.file_analysis.docx_parts import pseudonymize_docx_parts
from noirdoc.file_analysis.models import FileBlock
from noirdoc.file_analysis.reconstruction import reconstruct
from noirdoc.file_analysis.xlsx_parts import PartsResult
from noirdoc.file_reidentification.service import reidentify_file_bytes
from noirdoc.pseudonymization.engine import PseudonymizationEngine
from noirdoc.pseudonymization.mapper import PseudonymMapper
from tests.file_analysis.docx_helpers import (
    DOCX_MIME,
    clean_document_bytes,
    document_bytes,
    inject_app_props,
    inject_custom_props,
    inject_custom_xml,
    inject_people,
    inject_thumbnail,
)
from tests.file_analysis.xlsx_helpers import (
    SubstringDetector,
    assert_no_part_contains,
    part_names,
    read_part,
    rewrite_zip,
)


def _doc() -> DocumentObject:
    doc = Document()
    doc.add_paragraph("Zeile ohne PII")
    return doc


async def _redact(
    data: bytes,
    table: dict[str, str] | None = None,
    mapper: PseudonymMapper | None = None,
    *,
    apply: bool = True,
) -> tuple[bytes | None, PartsResult, PseudonymMapper]:
    mapper = mapper if mapper is not None else PseudonymMapper()
    new_bytes, result = await pseudonymize_docx_parts(
        data, SubstringDetector(table or {}), mapper, "de", apply=apply
    )
    return new_bytes, result, mapper


# ── docProps/core.xml ────────────────────────────────────


async def test_core_author_fields_forced():
    doc = _doc()
    doc.core_properties.author = "Anna Mueller"
    doc.core_properties.last_modified_by = "dora.klein@example.com"
    data = clean_document_bytes(doc)

    new_bytes, result, mapper = await _redact(data)

    assert new_bytes is not None
    out = Document(io.BytesIO(new_bytes))
    assert out.core_properties.author == "<<PERSON_1>>"
    assert out.core_properties.last_modified_by == "<<EMAIL_1>>"
    assert result.entity_types == {"PERSON": 1, "EMAIL": 1}
    assert result.classifications["core.author"] == "PERSON (forced)"
    assert result.classifications["core.last_modified_by"] == "EMAIL (forced)"
    assert mapper.reverse_lookup("<<PERSON_1>>") == "Anna Mueller"
    assert_no_part_contains(new_bytes, ["Anna Mueller", "dora.klein@example.com"])


async def test_core_free_text_fields_go_through_detector():
    doc = _doc()
    props = doc.core_properties
    props.title = "Akte Anna Mueller"
    props.subject = "Mandant Schmidt GmbH"
    props.comments = "Kontakt: anna@example.com"  # dc:description
    props.keywords = "Anna Mueller, Bestand"
    props.category = "Kanzlei"
    props.content_status = "Freigegeben von Anna Mueller"
    props.identifier = "Akte Anna Mueller"
    data = clean_document_bytes(doc)
    table = {
        "Anna Mueller": "PERSON",
        "Schmidt GmbH": "ORGANIZATION",
        "anna@example.com": "EMAIL",
    }

    new_bytes, result, _mapper = await _redact(data, table)

    assert new_bytes is not None
    out = Document(io.BytesIO(new_bytes)).core_properties
    assert out.title == "Akte <<PERSON_1>>"
    assert out.subject == "Mandant <<ORGANIZATION_1>>"
    assert out.comments == "Kontakt: <<EMAIL_1>>"
    assert out.keywords == "<<PERSON_1>>, Bestand"
    assert out.category == "Kanzlei"
    assert out.content_status == "Freigegeben von <<PERSON_1>>"
    assert out.identifier == "Akte <<PERSON_1>>"
    # title + keywords + content_status + identifier
    assert result.entity_types == {"PERSON": 4, "ORGANIZATION": 1, "EMAIL": 1}
    assert result.classifications["core.title"] == "PERSON (detected)"
    assert_no_part_contains(new_bytes, list(table))


async def test_default_template_noise_is_reported_not_leaked():
    """python-docx's own template ships creator, thumbnail and a customXml island."""
    data = document_bytes(_doc())  # deliberately NOT cleaned

    new_bytes, result, _mapper = await _redact(data)

    assert new_bytes is not None
    assert result.entity_types == {"PERSON": 1, "THUMBNAIL": 1, "CUSTOM_XML": 1}
    out = Document(io.BytesIO(new_bytes))
    assert out.core_properties.author == "<<PERSON_1>>"  # "python-docx" — no allowlist
    names = part_names(new_bytes)
    assert "docProps/thumbnail.jpeg" not in names
    assert not any(n.startswith("customXml/") for n in names)


# ── comment authors / initials ───────────────────────────


async def test_comment_author_and_initials_pseudonymized():
    doc = _doc()
    para = doc.add_paragraph("Siehe Anmerkung")
    doc.add_comment(para.runs, text="bitte pruefen", author="Dora Klein", initials="DK")
    data = clean_document_bytes(doc)

    new_bytes, result, mapper = await _redact(data)

    assert new_bytes is not None
    comment = next(iter(Document(io.BytesIO(new_bytes)).comments))
    assert comment.author == "<<PERSON_1>>"
    assert comment.initials == "<<PERSON_2>>"
    assert comment.text == "bitte pruefen"  # comment *text* is the body pass's job
    cid = comment.comment_id
    assert result.classifications[f"comment{cid}.author"] == "PERSON (forced)"
    # Initials duplicate the author identity: mapped but not counted twice.
    assert result.entity_types == {"PERSON": 1}
    assert mapper.reverse_lookup("<<PERSON_1>>") == "Dora Klein"
    assert mapper.reverse_lookup("<<PERSON_2>>") == "DK"
    assert_no_part_contains(new_bytes, ["Dora Klein"])


# ── docProps/app.xml ─────────────────────────────────────


async def test_app_props_pseudonymized_reversibly():
    data = inject_app_props(
        clean_document_bytes(_doc()),
        manager="Dora Klein",
        company="Schmidt GmbH",
        hyperlink_base="https://intranet.schmidt-gmbh.example/dora.klein/",
        titles_of_parts=["Akte Anna Mueller"],
    )
    table = {
        "Anna Mueller": "PERSON",
        "https://intranet.schmidt-gmbh.example/dora.klein/": "URL",
    }

    new_bytes, result, mapper = await _redact(data, table)

    assert new_bytes is not None
    assert result.classifications["app.Manager"] == "PERSON (forced)"
    assert result.classifications["app.Company"] == "ORGANIZATION (forced)"
    assert result.classifications["app.HyperlinkBase"] == "URL (detected)"
    assert result.classifications["app.TitlesOfParts.1"] == "PERSON (detected)"
    assert result.entity_types == {"PERSON": 2, "ORGANIZATION": 1, "URL": 1}
    # Reversible, unlike the XLSX side where openpyxl drops app.xml on write.
    assert mapper.reverse_lookup(mapper.lookup("Dora Klein") or "") == "Dora Klein"
    assert mapper.reverse_lookup(mapper.lookup("Schmidt GmbH") or "") == "Schmidt GmbH"
    app = read_part(new_bytes, "docProps/app.xml").decode()
    assert "&lt;&lt;PERSON_" in app
    assert_no_part_contains(
        new_bytes, ["Dora Klein", "Schmidt GmbH", "Anna Mueller", "intranet.schmidt-gmbh"]
    )


# ── docProps/custom.xml ──────────────────────────────────


async def test_custom_string_props_pseudonymized_other_types_untouched():
    data = inject_custom_props(
        clean_document_bytes(_doc()), {"Mandant": "Akte Schmidt GmbH", "AktenNr": 42}
    )

    new_bytes, result, _mapper = await _redact(data, {"Schmidt GmbH": "ORGANIZATION"})

    assert new_bytes is not None
    custom = read_part(new_bytes, "docProps/custom.xml").decode()
    assert "Schmidt GmbH" not in custom
    assert "&lt;&lt;ORGANIZATION_1&gt;&gt;" in custom
    assert "<vt:i4>42</vt:i4>" in custom
    assert result.classifications["custom.Mandant"] == "ORGANIZATION (detected)"
    assert result.entity_types == {"ORGANIZATION": 1}


# ── word/people.xml ──────────────────────────────────────


async def test_people_scrubbed_consistent_with_comment_author():
    doc = _doc()
    para = doc.add_paragraph("Siehe Anmerkung")
    doc.add_comment(para.runs, text="bitte pruefen", author="Dora Klein", initials="DK")
    data = inject_people(
        clean_document_bytes(doc),
        author="Dora Klein",
        user_id="S::dora.klein@example.com::12345",
    )

    new_bytes, result, mapper = await _redact(data)

    assert new_bytes is not None
    people = read_part(new_bytes, "word/people.xml").decode()
    assert "Dora Klein" not in people
    assert "dora.klein@example.com" not in people
    # Same identity, same placeholder as the comment's w:author.
    author_placeholder = mapper.lookup("Dora Klein")
    assert author_placeholder is not None
    assert author_placeholder.replace("<", "&lt;").replace(">", "&gt;") in people
    assert result.classifications["people.person1.author"] == "PERSON (forced)"
    assert result.classifications["people.person1.userId"] == "EMAIL (forced)"
    assert result.entity_types == {"PERSON": 2, "EMAIL": 1}
    assert_no_part_contains(new_bytes, ["Dora Klein", "dora.klein@example.com"])


# ── dropped parts ────────────────────────────────────────


async def test_thumbnail_dropped_and_counted():
    data = inject_thumbnail(clean_document_bytes(_doc()))

    new_bytes, result, _mapper = await _redact(data)

    assert new_bytes is not None
    assert "docProps/thumbnail.jpeg" not in part_names(new_bytes)
    assert b"thumbnail" not in read_part(new_bytes, "_rels/.rels")
    assert result.entity_types == {"THUMBNAIL": 1}
    assert result.classifications["docProps.thumbnail"] == "THUMBNAIL (dropped)"


async def test_custom_xml_island_dropped_and_counted():
    data = inject_custom_xml(
        clean_document_bytes(_doc()),
        xml_text="<kunde><name>Anna Mueller</name><iban>DE02120300000000202051</iban></kunde>",
    )

    new_bytes, result, _mapper = await _redact(data)

    assert new_bytes is not None
    assert not any(n.startswith("customXml/") for n in part_names(new_bytes))
    assert b"customXml" not in read_part(new_bytes, "word/_rels/document.xml.rels")
    assert result.entity_types == {"CUSTOM_XML": 1}
    assert result.classifications["customXml.item1"] == "CUSTOM_XML (dropped)"
    assert_no_part_contains(new_bytes, ["Anna Mueller", "DE02120300000000202051"])


async def test_unreadable_people_part_dropped_not_passed_through():
    data = inject_people(
        clean_document_bytes(_doc()), author="Dora Klein", user_id="dora@example.com"
    )
    data = rewrite_zip(data, {"word/people.xml": b"<w15:people"})

    new_bytes, result, _mapper = await _redact(data)

    assert new_bytes is not None
    assert "word/people.xml" not in part_names(new_bytes)
    assert result.entity_types == {"UNREADABLE_PART": 1}
    assert result.classifications["word/people.xml"] == "UNREADABLE_PART (dropped)"


# ── counting-only mode ───────────────────────────────────


async def test_apply_false_counts_without_writing_or_mapping():
    doc = _doc()
    doc.core_properties.author = "Anna Mueller"
    para = doc.add_paragraph("Text")
    doc.add_comment(para.runs, text="x", author="Dora Klein", initials="DK")
    data = inject_thumbnail(
        inject_app_props(clean_document_bytes(doc), manager="Emil Roth", company="Schmidt GmbH")
    )

    new_bytes, result, mapper = await _redact(data, apply=False)

    assert new_bytes is None
    # author + comment author + Manager + Company + thumbnail
    assert result.entity_types == {"PERSON": 3, "ORGANIZATION": 1, "THUMBNAIL": 1}
    assert mapper.get_mapping_summary() == {}


# ── reveal ───────────────────────────────────────────────


async def test_full_round_trip_restores_every_reversible_surface():
    doc = Document()
    doc.add_paragraph("Vertrag mit Anna Mueller")
    doc.sections[0].header.paragraphs[0].text = "Vertraulich: Anna Mueller"
    doc.sections[0].footer.paragraphs[0].text = "Kontakt: anna@example.com"
    para = doc.add_paragraph("Siehe Anmerkung")
    doc.add_comment(para.runs, text="Ruf Anna Mueller an", author="Dora Klein", initials="DK")
    props = doc.core_properties
    props.author = "Anna Mueller"
    props.title = "Akte Anna Mueller"
    data = clean_document_bytes(doc)
    data = inject_app_props(data, manager="Emil Roth", company="Schmidt GmbH")
    data = inject_custom_props(data, {"Mandant": "Schmidt GmbH"})
    data = inject_people(data, author="Dora Klein", user_id="S::dora.klein@example.com::1")
    data = inject_thumbnail(data)
    data = inject_custom_xml(data, xml_text="<kunde><name>Anna Mueller</name></kunde>")
    originals = [
        "Anna Mueller",
        "Dora Klein",
        "Emil Roth",
        "Schmidt GmbH",
        "anna@example.com",
        "dora.klein@example.com",
    ]
    table = {
        "Anna Mueller": "PERSON",
        "Dora Klein": "PERSON",
        "anna@example.com": "EMAIL",
        "Schmidt GmbH": "ORGANIZATION",
    }
    detector = SubstringDetector(table)
    mapper = PseudonymMapper()

    # The body pass, exactly as pipeline/sdk run it before the parts pass.
    from noirdoc.file_analysis.extractors.docx_ext import extract_docx

    block = FileBlock(
        content_bytes=data, mime_type=DOCX_MIME, source_path="test", source_type="file"
    )
    block.extracted_text = extract_docx(data)
    block.entities = await detector.detect(block.extracted_text)
    block.pseudonymized_text = PseudonymizationEngine().pseudonymize(
        block.extracted_text, block.entities, mapper
    )
    body_bytes = reconstruct(block)
    assert body_bytes is not None

    redacted, _result = await pseudonymize_docx_parts(body_bytes, detector, mapper, "de")
    assert redacted is not None
    assert_no_part_contains(redacted, originals)

    revealed = reidentify_file_bytes(redacted, DOCX_MIME, mapper.get_mapping_summary())
    assert revealed is not None
    out = Document(io.BytesIO(revealed))
    assert out.paragraphs[0].text == "Vertrag mit Anna Mueller"
    assert out.sections[0].header.paragraphs[0].text == "Vertraulich: Anna Mueller"
    assert out.sections[0].footer.paragraphs[0].text == "Kontakt: anna@example.com"
    comment = next(iter(out.comments))
    assert comment.author == "Dora Klein"
    assert comment.initials == "DK"
    assert comment.text == "Ruf Anna Mueller an"
    assert out.core_properties.author == "Anna Mueller"
    assert out.core_properties.title == "Akte Anna Mueller"
    app = read_part(revealed, "docProps/app.xml").decode()
    assert "Emil Roth" in app and "Schmidt GmbH" in app
    assert "Schmidt GmbH" in read_part(revealed, "docProps/custom.xml").decode()
    people = read_part(revealed, "word/people.xml").decode()
    assert "Dora Klein" in people and "dora.klein@example.com" in people
    # Dropped parts stay dropped — there is nothing truthful to restore.
    names = part_names(revealed)
    assert "docProps/thumbnail.jpeg" not in names
    assert not any(n.startswith("customXml/") for n in names)


def test_reveal_placeholder_only_in_metadata_still_rewrites():
    """Regression: _reidentify_docx returned the input unchanged when no *body* run changed."""
    doc = _doc()
    doc.core_properties.author = "<<PERSON_1>>"
    data = clean_document_bytes(doc)

    revealed = reidentify_file_bytes(data, DOCX_MIME, {"<<PERSON_1>>": "Anna Mueller"})

    assert revealed is not None
    assert revealed != data
    assert Document(io.BytesIO(revealed)).core_properties.author == "Anna Mueller"


def test_reveal_covers_headers_footers_and_comment_text():
    """Regression: the reveal walker only visited body paragraphs and tables."""
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Von <<PERSON_1>>"
    doc.sections[0].footer.paragraphs[0].text = "An <<PERSON_1>>"
    para = doc.add_paragraph("Text")
    doc.add_comment(para.runs, text="Frag <<PERSON_1>>", author="<<PERSON_2>>", initials="DK")
    data = clean_document_bytes(doc)

    revealed = reidentify_file_bytes(
        data, DOCX_MIME, {"<<PERSON_1>>": "Anna Mueller", "<<PERSON_2>>": "Dora Klein"}
    )

    assert revealed is not None
    out = Document(io.BytesIO(revealed))
    assert out.sections[0].header.paragraphs[0].text == "Von Anna Mueller"
    assert out.sections[0].footer.paragraphs[0].text == "An Anna Mueller"
    comment = next(iter(out.comments))
    assert comment.text == "Frag Anna Mueller"
    assert comment.author == "Dora Klein"

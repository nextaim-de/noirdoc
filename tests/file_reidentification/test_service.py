from __future__ import annotations

import io
from typing import TYPE_CHECKING

from noirdoc.file_reidentification.service import reidentify_file_bytes

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

MAPPINGS = {
    "<<PERSON_1>>": "Max Müller",
    "<<EMAIL_1>>": "max@test.de",
}


# ── Plain text ───────────────────────────────────────────


def test_reidentify_plain_text():
    content = b"Budget for <<PERSON_1>> (<<EMAIL_1>>)"
    result = reidentify_file_bytes(content, "text/plain", MAPPINGS)
    assert result is not None
    assert result == "Budget for Max Müller (max@test.de)".encode()


def test_reidentify_csv():
    content = b"Name,Email\n<<PERSON_1>>,<<EMAIL_1>>"
    result = reidentify_file_bytes(content, "text/csv", MAPPINGS)
    assert result is not None
    text = result.decode()
    assert "Max Müller" in text
    assert "max@test.de" in text


def test_reidentify_text_with_charset():
    content = b"Hello <<PERSON_1>>."
    result = reidentify_file_bytes(content, "text/plain; charset=utf-8", MAPPINGS)
    assert result is not None
    assert b"Max" in result


# ── DOCX ─────────────────────────────────────────────────


def test_reidentify_docx():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Report for <<PERSON_1>>")
    buf = io.BytesIO()
    doc.save(buf)

    result = reidentify_file_bytes(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        MAPPINGS,
    )
    assert result is not None

    # Verify the result is a valid DOCX with reidentified text
    result_doc = Document(io.BytesIO(result))
    text = result_doc.paragraphs[0].text
    assert "Max Müller" in text
    assert "<<PERSON_1>>" not in text


def test_reidentify_docx_table():
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "<<PERSON_1>>"
    buf = io.BytesIO()
    doc.save(buf)

    result = reidentify_file_bytes(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        MAPPINGS,
    )
    assert result is not None

    result_doc = Document(io.BytesIO(result))
    assert result_doc.tables[0].rows[0].cells[1].text == "Max Müller"


def _add_hyperlink(para: Paragraph, text: str, url: str) -> None:
    """Append a real ``w:hyperlink`` (with one run) to *para*."""
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    r_id = para.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


def test_reidentify_docx_pseudonym_inside_hyperlink():
    """Reveal round-trip: a pseudonym in hyperlink text must be reidentified in place."""
    from docx import Document

    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Kontakt: ")
    _add_hyperlink(para, "<<EMAIL_1>>", "mailto:masked@example.invalid")
    buf = io.BytesIO()
    doc.save(buf)

    result = reidentify_file_bytes(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        MAPPINGS,
    )
    assert result is not None

    out = Document(io.BytesIO(result)).paragraphs[0]
    assert out.text == "Kontakt: max@test.de"
    # The original must land inside the hyperlink, not in the plain runs
    assert out.hyperlinks[0].text == "max@test.de"
    assert "max@test.de" not in "".join(run.text for run in out.runs)


def test_reidentify_docx_keeps_unrelated_hyperlink_intact():
    """A hyperlink next to a revealed pseudonym must not be duplicated into runs."""
    from docx import Document

    doc = Document()
    para = doc.add_paragraph()
    para.add_run("<<PERSON_1>> (siehe ")
    _add_hyperlink(para, "www.firma.de", "https://www.firma.de")
    para.add_run(")")
    buf = io.BytesIO()
    doc.save(buf)

    result = reidentify_file_bytes(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        MAPPINGS,
    )
    assert result is not None

    out = Document(io.BytesIO(result)).paragraphs[0]
    assert out.text == "Max Müller (siehe www.firma.de)"
    assert out.hyperlinks[0].text == "www.firma.de"
    assert "www.firma.de" not in "".join(run.text for run in out.runs)


def test_reidentify_docx_no_pseudonyms():
    from docx import Document

    doc = Document()
    doc.add_paragraph("No pseudonyms here.")
    buf = io.BytesIO()
    doc.save(buf)

    original = buf.getvalue()
    result = reidentify_file_bytes(
        original,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        MAPPINGS,
    )
    # Returns original bytes when nothing changed
    assert result == original


# ── XLSX ─────────────────────────────────────────────────


def test_reidentify_xlsx():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Name"
    ws["B1"] = "<<PERSON_1>>"
    ws["A2"] = "Email"
    ws["B2"] = "<<EMAIL_1>>"
    buf = io.BytesIO()
    wb.save(buf)

    result = reidentify_file_bytes(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        MAPPINGS,
    )
    assert result is not None

    from openpyxl import load_workbook

    result_wb = load_workbook(io.BytesIO(result))
    ws = result_wb.active
    assert ws["B1"].value == "Max Müller"
    assert ws["B2"].value == "max@test.de"


def test_reidentify_xlsx_no_pseudonyms():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Hello"
    buf = io.BytesIO()
    wb.save(buf)

    original = buf.getvalue()
    result = reidentify_file_bytes(
        original,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        MAPPINGS,
    )
    assert result == original


# ── Unsupported formats ──────────────────────────────────


def test_unsupported_pdf():
    result = reidentify_file_bytes(b"fake pdf", "application/pdf", MAPPINGS)
    assert result is None


def test_unsupported_pptx():
    result = reidentify_file_bytes(
        b"fake pptx",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        MAPPINGS,
    )
    assert result is None


def test_unsupported_image():
    result = reidentify_file_bytes(b"fake png", "image/png", MAPPINGS)
    assert result is None


def test_empty_mappings():
    result = reidentify_file_bytes(b"hello", "text/plain", {})
    assert result is None
